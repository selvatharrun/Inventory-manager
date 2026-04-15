"""Polish Final Report.docx content, structure text, and formatting."""

from __future__ import annotations

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _heading_indices(doc: Document) -> List[int]:
    indices: List[int] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if p.style.name.startswith("Heading") or t.startswith("CHAPTER") or t in {"ABSTRACT", "REFERENCES", "WORKLOG"}:
            indices.append(i)
    return sorted(indices)


def _find_index(doc: Document, marker: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == marker:
            return i
    raise ValueError(f"Marker not found: {marker}")


def _has_marker(doc: Document, marker: str) -> bool:
    """Return True when a paragraph exactly matches marker text."""
    for p in doc.paragraphs:
        if p.text.strip() == marker:
            return True
    return False


def _replace_section(doc: Document, marker: str, new_lines: List[str]) -> None:
    start = _find_index(doc, marker)
    heading_idxs = _heading_indices(doc)
    next_candidates = [idx for idx in heading_idxs if idx > start]
    end = next_candidates[0] if next_candidates else len(doc.paragraphs)

    body_indices = [idx for idx in range(start + 1, end)]
    for idx, line in zip(body_indices, new_lines):
        doc.paragraphs[idx].text = line
    for idx in body_indices[len(new_lines) :]:
        doc.paragraphs[idx].text = ""


def _polish_content(doc: Document) -> None:
    replacements: Dict[str, List[str]] = {
        "ABSTRACT": [
            "This project addresses operational challenges in inventory planning where manual spreadsheet-based decisions are slow, inconsistent, and difficult to scale.",
            "The developed Inventory Optimization AI Agent analyzes SKU-level stock records and generates advisory recommendations for replenishment, urgency, and risk control.",
            "The implementation uses a local-first architecture: LangGraph for orchestration, FastMCP for tool interfaces, a hybrid knowledge graph (Neo4j with NetworkX fallback), and Ollama for batched plain-English explanations.",
            "A synthetic dataset of 25 SKUs was used to validate status classification, context enrichment, fallback behavior, and output schema compliance.",
            "Evaluation focused on reliability, explanation completeness, and end-to-end latency under CPU-only constraints on a standard laptop.",
        ],
        "1.1 Background of the Study": [
            "Modern supply chains require rapid, evidence-based inventory decisions to avoid stockouts, excess holding costs, and missed service targets.",
            "Many organizations still depend on manually maintained spreadsheets and fragmented rules, which creates inconsistency across planners and delays response to demand changes.",
            "Recent advances in AI and workflow orchestration enable practical decision-support systems that combine deterministic metrics with explainable language outputs.",
            "This study designs a local, advisory-only inventory agent that helps planners prioritize actions while preserving human oversight.",
        ],
        "1.2 Problem Statement": [
            "Inventory teams often face decision friction due to siloed data, inconsistent threshold logic, and limited contextual visibility across SKU categories and seasonal effects.",
            "Traditional approaches either require technical scripting capability or rely on cloud-heavy automation that may not satisfy latency, privacy, or deployment constraints.",
            "The project therefore targets a local, CPU-friendly decision-support workflow that produces standardized and explainable recommendations without automating procurement actions.",
        ],
        "1.3 Objectives of the Study": [
            "The core objective is to implement a production-style prototype of an Inventory Optimization AI Agent for advisory decision support.",
            "Specific objectives:",
            "Build validated data ingestion for CSV/JSON inventory inputs.",
            "Compute SKU-level metrics such as days-of-stock, reorder quantity, urgency, and velocity trend.",
            "Integrate hybrid context enrichment using Neo4j with deterministic NetworkX fallback.",
            "Generate batched plain-English explanations with strict timeout and template fallback.",
            "Enforce schema-validated outputs with mandatory advisory disclaimer.",
        ],
        "1.4 Scope and Limitations": [
            "Scope:",
            "Local execution on standard laptops (CPU-only, 16GB RAM target profile).",
            "Structured inventory inputs with configurable thresholds and scenario overrides.",
            "JSON output for downstream reporting and optional Streamlit display.",
            "Limitations:",
            "Recommendations are advisory and do not trigger order placement.",
            "Quality of recommendations depends on input data quality and threshold configuration.",
            "Complex multi-echelon optimization is out of scope for the current implementation.",
        ],
        "1.5 Significance of the Study": [
            "The project demonstrates that practical AI decision support can run locally without high-end infrastructure.",
            "It improves consistency in inventory planning by standardizing metric computation, fallback handling, and explanation format.",
            "The architecture is modular and testable, making it suitable for iterative enhancement in enterprise or academic settings.",
        ],
        "2.1 Introduction to the Literature": [
            "Prior work in inventory analytics spans deterministic EOQ/safety-stock methods, statistical forecasting pipelines, and ML-enhanced demand modeling.",
            "Recent research extends beyond pure optimization by emphasizing explainability, human-in-the-loop controls, and resilient orchestration across data and model services.",
            "This review frames the current project as a hybrid approach: rule-grounded metrics, graph context enrichment, and constrained LLM explanation generation.",
        ],
        "2.2 Related Work": [
            "Classical inventory control models provide strong mathematical foundations but often require manual interpretation and frequent recalibration in dynamic environments.",
            "Business intelligence tooling improves visibility but typically lacks automated advisory narratives and robust fallback orchestration.",
            "LLM-enabled systems improve accessibility through natural-language outputs, yet many implementations depend heavily on cloud APIs and per-item calls that increase latency and cost.",
            "Hybrid pipelines that combine deterministic rules with constrained language generation are emerging as a practical middle ground for operational decision support.",
        ],
        "Key Challenges of Existing Systems": [
            "Limited explainability for non-technical stakeholders.",
            "Weak fallback strategies when external model or graph services fail.",
            "High latency from per-record model calls under CPU constraints.",
            "Inconsistent output contracts across tools and orchestration layers.",
            "Insufficient local-first implementations for privacy-sensitive operations.",
        ],
        "2.3 Research Gaps": [
            "Existing solutions rarely combine local execution, strict schema validation, and deterministic fallback guarantees in one pipeline.",
            "Few implementations report practical constraints such as CPU-only inference and sub-5-second response targets for batch SKU analysis.",
            "This project addresses those gaps through batched LLM calls, cache-aware graph enrichment, and test-driven fallback validation.",
        ],
        "3.1 Introduction": [
            "This chapter presents the end-to-end methodology used to design and implement the Inventory Optimization AI Agent.",
            "The proposed approach combines deterministic inventory formulas, rule-based policy checks, graph context retrieval, and constrained language explanations.",
            "The architecture prioritizes reliability and low-latency local execution over autonomous control.",
        ],
        "3.2 Proposed System / Methodology": [
            "The system follows a staged pipeline from validated data ingestion to schema-checked recommendations.",
            "Input records are loaded and normalized, per-SKU metrics are computed, contextual enrichment is attached, and rules are applied before final formatting.",
            "A single batched LLM request is used for explanation generation; when unavailable or slow, deterministic templates guarantee complete output coverage.",
        ],
        "Workflow Pipeline": [
            "Load Data -> Calculate Metrics -> Enrich Context -> Apply Rules -> Generate Recs -> Explain (Batch LLM) -> Template Fallback -> Format -> Validate",
            "This pipeline is implemented in LangGraph with node-level state transitions and fault-tolerant fallback behavior.",
        ],
        "Key Features of the Proposed System": [
            "Local-first execution with advisory-only decision outputs.",
            "Hybrid graph context retrieval: cache -> Neo4j -> NetworkX fallback.",
            "Single-call batched LLM explanation strategy for latency control.",
            "Schema validation with mandatory disclaimer and non-empty explanation checks.",
        ],
        "3.3 Overall Workflow Diagram": [
            "The workflow separates deterministic decision logic from probabilistic explanation generation to preserve reliability under partial service failure.",
        ],
        "Figure 3.1: System Workflow": [
            "Components:",
            "CLI/Optional Streamlit UI",
            "FastMCP Tool Layer",
            "LangGraph Orchestrator",
            "Knowledge Context Layer (Neo4j + NetworkX)",
            "Local LLM Layer (Ollama)",
        ],
        "Explanation:": [
            "Inventory records are validated and transformed into SKU-level metrics.",
            "Context is retrieved through cache-aware graph lookup.",
            "Rules and metrics are merged into recommendation payloads.",
            "Explanations are generated in one batched call or filled by deterministic templates.",
            "Final output is schema-validated and emitted as JSON plus human-readable summary.",
        ],
        "3.4 Module Description": [
            "The implementation is modular, with each module independently testable through MCP tool contracts or graph-node execution.",
        ],
        "1. User Interface Module": [
            "This module provides the execution entry points for planners.",
            "Functions:",
            "Run analysis through CLI flags and scenario overrides.",
            "Optional Streamlit dashboard for upload, status cards, and report review.",
            "Outputs:",
            "Human-readable summary, structured JSON, and markdown report card.",
        ],
        "2. Backend Processing Module": [
            "This module coordinates loading, validation, and orchestration.",
            "Functions:",
            "Input normalization and required-field validation.",
            "LangGraph state execution and node transitions.",
            "Failure capture via warnings/errors in metadata.",
        ],
        "3. AI Model Module": [
            "This module generates plain-English advisory explanations.",
            "Runtime:",
            "Ollama with small local model (llama3.2:1b or equivalent).",
            "Controls:",
            "Single batched call, low temperature, JSON mode, and 4-second timeout.",
            "Fallback:",
            "Deterministic templates ensure explanation completeness when model calls fail.",
        ],
        "4. Visualization Module": [
            "This module handles output presentation rather than chart rendering.",
            "Functions:",
            "Generate recommendation records with status, metrics, context, and explanations.",
            "Render table/JSON output and optional UI cards.",
            "Produce markdown report card for executive review.",
        ],
        "5. Data Processing Module": [
            "This module guarantees data quality before decision computation.",
            "Functions:",
            "Type normalization and invalid-row isolation.",
            "Metric computation per SKU with threshold-based status classification.",
            "Scenario override application for what-if comparisons.",
        ],
        "3.5 Algorithms / Models Used": [
            "The system combines deterministic formulas with constrained LLM post-processing.",
        ],
        "1. Natural Language Processing (NLP)": [
            "NLP is used only for explanation generation, not for core decision arithmetic.",
            "This separation improves traceability and prevents opaque metric outcomes.",
        ],
        "2. Large Language Models (LLMs)": [
            "A compact local model is used to generate short, structured explanations for all SKUs in one call.",
            "Output is constrained to JSON with fixed fields: explanation, action, confidence.",
        ],
        "3. Visualization Mapping Algorithm": [
            "Decision mapping logic translates metric status and urgency into recommended actions.",
            "Rule IDs are assigned by status bands and merged with graph-derived risk tags.",
        ],
        "4. Plotly JSON Generation": [
            "The final payload is emitted as schema-validated JSON for downstream consumption.",
            "Each recommendation includes SKU metrics, context source, action text, confidence, and disclaimer compliance.",
        ],
        "3.6 Summary": [
            "This chapter established the methodology for a reliable local decision-support pipeline.",
            "The design emphasizes modularity, fallback resilience, and verifiable outputs under constrained hardware.",
        ],
        "4.1 Introduction": [
            "This chapter describes the practical implementation of the proposed architecture and runtime behavior.",
            "The implementation is Python-based and structured into tools, graph nodes, knowledge adapters, and interface layers.",
        ],
        "Hardware Requirements": [
            "Processor: Intel i5 / AMD Ryzen 5 or equivalent",
            "RAM: 16 GB recommended (8 GB minimum)",
            "Storage: 10 GB free disk space",
            "GPU: Not required",
        ],
        "Software Requirements": [
            "Operating System: Windows / macOS / Linux",
            "Language Runtime: Python 3.11+",
            "Core Stack: LangGraph, FastMCP, Pydantic, jsonschema",
            "Knowledge Layer: Neo4j (optional) + NetworkX fallback",
            "Model Runtime: Ollama local inference",
            "Testing: Pytest",
        ],
        "4.3 Tools / Libraries Used": [
            "The system uses a focused Python stack optimized for local reliability and testability.",
        ],
        "1. Next.js": [
            "LangGraph: Orchestrates stateful workflow execution across deterministic and model-driven nodes.",
        ],
        "2. React.js": [
            "FastMCP: Exposes independently testable tools for data loading, metrics, rules, and graph queries.",
        ],
        "3. Tailwind CSS": [
            "Diskcache: Provides TTL-based caching for graph lookups to reduce repeated latency.",
        ],
        "5. Qwen 2.5 Coder Model": [
            "Local Ollama Model (llama3.2:1b): Generates batched explanations in JSON mode with low temperature.",
        ],
        "6. Plotly.js": [
            "NetworkX + Neo4j: Hybrid context layer for seasonal and category intelligence with fallback safety.",
        ],
        "7. TypeScript": [
            "Pytest + Jsonschema: Enforce behavior correctness and output contract validation.",
        ],
        "4.4 Implementation Details": [
            "Implementation follows a deterministic-first execution model with constrained AI post-processing.",
        ],
        "Step 1: Dataset Upload": [
            "Inventory CSV/JSON is provided via CLI or optional Streamlit interface.",
            "Input path is validated before pipeline execution.",
        ],
        "Step 2: Data Preprocessing": [
            "Rows are normalized to required numeric/text schema.",
            "Invalid rows are skipped and recorded in warnings metadata.",
        ],
        "Step 3: User Query Processing": [
            "A run configuration is assembled with thresholds and optional scenario overrides.",
            "No natural-language query is required for core metric computation.",
        ],
        "Step 4: AI Model Integration": [
            "LLM is called once with batched SKU payloads for explanation generation.",
            "Timeout and failure conditions route output to deterministic templates.",
        ],
        "Step 5: Visualization Configuration Generation": [
            "Per-SKU advisory records are assembled with status, urgency, rule matches, and context annotations.",
        ],
        "Step 6: Chart Rendering": [
            "Results are rendered as CLI table/JSON and optional Streamlit cards.",
        ],
        "Step 7: Output Display": [
            "Validated JSON and markdown report are written to the results directory.",
            "Every output includes the mandatory advisory disclaimer.",
        ],
        "4.5 Dataset Collection and Descriptions": [
            "The implementation uses a synthetic but realistic SKU inventory dataset for repeatable local testing.",
        ],
        "Dataset Characteristics": [
            "Format: CSV and JSON",
            "Size: 25 SKUs in baseline dataset",
            "Fields: stock, velocity, lead time, safety stock, category",
            "Distribution: healthy/watch/critical/overstock coverage",
        ],
        "Dataset Processing Steps": [
            "File load and schema validation",
            "Type normalization and invalid-row filtering",
            "Metric derivation and status classification",
            "Context enrichment and rule mapping",
        ],
        "Example Dataset Use Cases": [
            "Daily replenishment review",
            "Category-level risk monitoring",
            "Lead-time what-if scenario comparison",
            "Fallback reliability verification",
        ],
        "Data Preprocessing Techniques": [
            "Missing-field validation",
            "Numeric type coercion",
            "Invalid record isolation",
            "Derived trend feature computation",
        ],
        "4.6 Pseudocode / Code Snippets": [
            "A simplified pseudocode representation of the implemented pipeline is provided below.",
        ],
        "Pseudocode": [
            "START",
            "LOAD config and inventory dataset",
            "VALIDATE required fields and normalize types",
            "FOR each SKU: compute days_of_stock, reorder_qty, urgency, trend, status",
            "ENRICH each SKU via cache -> Neo4j -> NetworkX fallback",
            "APPLY status rules and build recommendation payloads",
            "CALL LLM once with batched SKU summaries (JSON mode, timeout=4s)",
            "IF LLM fails or times out: fill explanations using templates",
            "FORMAT final JSON, validate schema, append disclaimer",
            "WRITE JSON + markdown report",
            "END",
        ],
        "Sample API Flow (Conceptual)": [
            "Input: inventory CSV/JSON + thresholds configuration",
            "Tool Calls: load_data -> calc_metrics -> query_graph -> fetch_rules",
            "Graph Flow: generate_recs -> explain_llm -> template_fallback -> format_output -> validate_output",
            "Output: advisory JSON report with per-SKU action, confidence, and explanation",
        ],
        "4.7 Summary": [
            "The implementation achieves a modular, testable, and local-first architecture aligned with project constraints.",
            "Fallback handling and schema checks ensure reliable outputs under partial component failure.",
        ],
        "5.1 Experimental Results": [
            "Experiments were executed on CPU-only local environment using the synthetic 25-SKU dataset.",
            "Results confirmed correct status classification, complete explanation coverage, and robust fallback transitions.",
        ],
        "Test Case 1: Sales Dataset": [
            "Case: Critical stock detection.",
            "Result: Low-coverage SKUs were correctly flagged as critical with positive reorder quantities.",
        ],
        "Test Case 2: Category Distribution": [
            "Case: Context enrichment under Neo4j-disabled mode.",
            "Result: NetworkX fallback supplied seasonal/category context and remained schema-compliant.",
        ],
        "Test Case 3: Regional Comparison": [
            "Case: Cache efficiency on repeated runs.",
            "Result: Subsequent graph lookups resolved from cache with reduced context-query overhead.",
        ],
        "Test Case 4: Invalid Query Handling": [
            "Case: Forced LLM unavailability/timeout.",
            "Result: Template fallback produced non-empty explanations for all SKUs.",
        ],
        "Observations": [
            "Metric outputs remained deterministic across repeated runs with identical inputs.",
            "Fallback order behaved as designed for graph and LLM layers.",
            "Output contract remained valid even under degraded component availability.",
        ],
        "Performance Metrics": [
            "Total run latency (CPU-only)",
            "Explanation completeness ratio",
            "Schema validation pass rate",
            "Fallback activation frequency",
        ],
        "Analysis": [
            "Batched explanation generation reduced overhead compared with per-SKU prompting.",
            "Template fallback maintained output continuity when model runtime was unavailable.",
            "End-to-end response remained within practical limits for small batch planning runs.",
        ],
        "5.3 Visualization of Results": [
            "Results are presented as structured recommendation records and summary indicators suitable for operational review.",
        ],
        "Types of Visualizations Generated": [
            "Critical SKU priority list",
            "Watch-list planning summary",
            "Overstock exposure snapshot",
            "Context-source/fallback status markers",
        ],
        "Graphical Representation": [
            "Figure 5.1: SKU status distribution",
            "Figure 5.2: Reorder urgency comparison",
            "Figure 5.3: Context source and fallback trace",
        ],
        "5.4 Interpretation and Analysis": [
            "The system demonstrates practical value for planner decision support by combining deterministic metrics with concise explanations.",
            "Human review remains central, while automation improves consistency and response speed.",
        ],
        "Key Insights": [
            "Deterministic metric logic provides stable decision baselines.",
            "Hybrid context improves prioritization fidelity.",
            "Batched LLM design supports low-latency explanation generation.",
        ],
        "Performance Strengths": [
            "Local execution with no cloud dependency requirement.",
            "Comprehensive fallback coverage across critical components.",
            "Strong testability at tool and integration levels.",
        ],
        "Behavior Analysis": [
            "System behavior is predictable under valid inputs and remains resilient under induced failures.",
            "Metadata warnings/errors provide clear diagnostics for debugging and monitoring.",
        ],
        "5.5 Limitations of Results": [
            "The prototype remains advisory and does not optimize multi-echelon replenishment schedules.",
            "Quality depends on threshold configuration and input cleanliness.",
            "Model explanation quality may vary with local runtime conditions.",
        ],
        "1. Dependency on LLM Accuracy": [
            "LLM outputs are constrained but still probabilistic; fallback templates mitigate this risk.",
        ],
        "2. Limited Handling of Complex Queries": [
            "The current version focuses on single-pass SKU advisory output rather than deep analytical dialogue.",
        ],
        "3. Dataset Constraints": [
            "Current validation assumes required schema fields are present and correctly mapped.",
        ],
        "5. Performance on Large Datasets": [
            "For very large SKU sets, chunking and incremental processing should be added in future iterations.",
        ],
        "6.1 Conclusion Drawn from the Results": [
            "The project successfully delivers a local, explainable, and testable Inventory Optimization AI Agent aligned with CPU-only operational constraints.",
            "The implemented architecture demonstrates reliable advisory generation through deterministic metrics, hybrid context enrichment, and robust fallback control.",
            "Test outcomes confirm schema consistency, explanation completeness, and practical run-time behavior for planning workflows.",
            "Overall, the system provides a strong foundation for safe human-in-the-loop inventory decision support.",
        ],
        "6.2 Future Work": [
            "Add chunked processing for larger SKU volumes and improved throughput.",
            "Introduce richer policy simulation for service-level and cost trade-off analysis.",
            "Expand graph context with supplier reliability and promotion events.",
            "Add calibration workflows for threshold tuning by category.",
            "Integrate audit dashboards for recommendation acceptance and outcome tracking.",
        ],
    }

    # Heading text corrections while preserving hierarchy.
    heading_text_swaps = {
        "4. Visualization Module": "4. Decision Output Module",
        "4. Plotly JSON Generation": "4. Structured JSON Generation",
        "5.3 Visualization of Results": "5.3 Presentation of Results",
    }

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in heading_text_swaps:
            p.text = heading_text_swaps[text]

    for marker, lines in replacements.items():
        if _has_marker(doc, marker):
            _replace_section(doc, marker, lines)

    # Update references section to project-relevant sources.
    if _has_marker(doc, "REFERENCES"):
        _replace_section(
            doc,
            "REFERENCES",
            [
                "LangChain. (2026). LangGraph documentation.",
                "Model Context Protocol. (2026). FastMCP and MCP tooling reference.",
                "Neo4j. (2026). Neo4j Graph Database documentation.",
                "Hagberg, A., Schult, D., & Swart, P. (2008). Exploring network structure using NetworkX.",
                "Ollama. (2026). Local model runtime documentation.",
                "JSON Schema Organization. (2026). JSON Schema specification.",
                "Pytest Development Team. (2026). Pytest documentation.",
                "DiskCache Maintainers. (2026). Diskcache Python library documentation.",
                "Pydantic Team. (2026). Pydantic validation framework documentation.",
                "Snyder, R. et al. (2012). Inventory management and demand forecasting: methods and practice.",
                "Silver, E., Pyke, D., & Thomas, D. (2016). Inventory and Production Management in Supply Chains.",
            ],
        )


def _apply_formatting(doc: Document) -> None:
    abstract_idx = _find_index(doc, "ABSTRACT")

    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue

        style = p.style.name
        if style == "Heading 1":
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif style == "Heading 2":
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif style == "Heading 3":
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        elif i >= abstract_idx:
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in p.runs:
            run.font.name = "Times New Roman"
            if style == "Heading 1":
                run.font.size = Pt(14)
                run.bold = True
            elif style == "Heading 2":
                run.font.size = Pt(13)
                run.bold = True
            elif style == "Heading 3":
                run.font.size = Pt(12)
                run.bold = True
            else:
                run.font.size = Pt(12)


def polish_docx(path: Path) -> None:
    doc = Document(str(path))
    _polish_content(doc)
    _apply_formatting(doc)
    try:
        doc.save(str(path))
    except PermissionError:
        fallback = path.with_name(f"{path.stem} - polished{path.suffix}")
        doc.save(str(fallback))
        print(f"Could not overwrite locked file. Saved polished copy to: {fallback}")


if __name__ == "__main__":
    polish_docx(Path("Final Report.docx"))
